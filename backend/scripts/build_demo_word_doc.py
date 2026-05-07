# -*- coding: utf-8 -*-
"""Generate 演示Demo-评委导读.docx from canonical outline (no external md parse)."""
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


def add_code_block(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int):
    doc.add_heading(text, level=level)


def add_shot_block(doc: Document, title: str, judge_text: str, bracket_instruction: str):
    """Judge-facing title + explanation; bracket line for presenter only."""
    add_heading(doc, title, level=2)
    doc.add_paragraph(judge_text)
    p = doc.add_paragraph()
    run = p.add_run(bracket_instruction)
    run.font.size = Pt(10)
    run.italic = True
    doc.add_paragraph()  # blank for paste area


def main():
    root = Path(__file__).resolve().parents[2]
    # ASCII 主文件名，避免部分 Windows/终端环境下中文路径乱码；另存中文副本供归档
    out = root / "Agentrust-Demo-Judges.docx"
    out_alias = root / "\u6f14\u793aDemo-\u8bc4\u59d4\u5bfc\u8bfb.docx"

    doc = Document()
    doc.add_heading("Agentrust 演示说明（评委导读）", 0)

    intro = (
        "本文说明演示所需的命令、运行顺序、终端输出含义，以及 Dashboard / Swagger 核验方式。"
        "评委可按章节对照运行结果。"
    )
    doc.add_paragraph(intro)

    # --- 1 ---
    add_heading(doc, "1. 演示目标", level=1)
    add_table(
        doc,
        ["能力", "演示如何体现"],
        [
            ["IAM 网关与健康检查", "一键脚本先启动 IAM，/health 探活通过后再继续"],
            ["多 Agent 并行服务", "企业数据 / 外部检索 / 文档助手各占独立端口"],
            ["正常业务链路", "demo_cycle4_normal.py：同一 task_id 下 IAM 多次 execute 成功"],
            ["越权与拦截", "demo_cycle4_abnormal.py：企业 Agent 403；IAM 能力与动作不匹配 403 并入审计"],
            ["审计与可视化", "控制台输出 task_id、session_token、agent_id；任务链路或 GET /audit/trace/{task_id} 可查"],
        ],
    )

    # --- 2 ---
    add_heading(doc, "2. 环境前置条件", level=1)
    for line in [
        "操作系统：Windows 10/11（脚本为 .bat；Linux/macOS 可参考 run_demo.sh）。",
        "Python 3.10+，已配置 pip。",
        "Node.js / npm：用于 Dashboard；若未安装则脚本跳过前端，后端演示仍可完成。",
        "请将下文路径中的 E:\\yy_project\\Agentrust 替换为你的本地工程路径。",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_paragraph("可选：首次拉代码后安装依赖并初始化数据库：")
    add_code_block(doc, r"cd /d E:\yy_project\Agentrust\backend" "\n" r"scripts\setup.bat")
    doc.add_paragraph(
        "setup.bat：进入 backend → 若无则创建 venv → pip install -r requirements.txt → "
        "创建 data → python scripts/init_db.py 初始化库表与 CA 根密钥。"
    )
    doc.add_paragraph(
        "若需清空数据库：执行 backend\\scripts\\reset.bat 后，再执行 setup.bat 或 init_db.py。"
    )

    # --- 3 ---
    add_heading(doc, "3. 端口与服务一览", level=1)
    add_table(
        doc,
        ["端口", "服务", "说明"],
        [
            ["8000", "IAM（FastAPI main:app）", "认证、resources/execute、审计 API"],
            ["8001", "企业数据 Agent", "演示：文档助手白名单、/read"],
            ["8002", "外部检索 Agent", "Mock 公开检索"],
            ["8003", "文档助手 Agent", "Mock 报告编排"],
            ["5173", "Dashboard（Vite）", "任务链路、审计列表"],
        ],
    )

    # --- 4 ---
    add_heading(doc, "4. 步骤一：一键演示（run_demo.bat）", level=1)

    add_heading(doc, "4.1 执行的命令", level=2)
    doc.add_paragraph("在 CMD 或 PowerShell 中执行（亦可双击运行）：")
    add_code_block(doc, r"cd /d E:\yy_project\Agentrust\backend\scripts" "\n" r"run_demo.bat")

    add_heading(doc, "4.2 运行顺序说明", level=2)
    steps = [
        "解析后端根目录（脚本所在目录的上一级 backend）。若存在 venv，优先使用该环境中的 Python。",
        "启动 IAM（端口 8000）：新开 CMD 窗口，在窗口内设置 UVICORN_APP、UVICORN_PORT 后调用 spawn_uvicorn.bat，"
        "避免 Windows 下将 main:app 当作路径语法解析出错，并避免多窗口环境变量互相覆盖。",
        "等待约 8 秒后调用 wait_health.py 探测 http://127.0.0.1:8000/health；失败则中止并提示排查。",
        "依次启动企业数据 Agent（8001）、外部检索（8002）、文档助手（8003）。",
        "若本机存在 npm：启动 Dashboard（5173）。",
        "依次执行 demo_cycle4_normal.py、demo_cycle4_abnormal.py。",
        "尝试打开浏览器 http://localhost:5173；主窗口 pause，便于阅读输出。",
        "仅启动服务不跑演示：可先执行 set SKIP_DEMOS=1 再运行 run_demo.bat。",
    ]
    for s in steps:
        doc.add_paragraph(s, style="List Number")

    add_heading(doc, "4.3 主窗口典型输出与含义", level=2)
    doc.add_paragraph("以下为应出现的输出类型（task_id 等每次运行不同）：")
    add_table(
        doc,
        ["输出片段", "含义"],
        [
            ["ROOT=…\\backend", "后端根目录识别正确"],
            ["Waiting for IAM …/health → OK", "IAM 已就绪"],
            ["Cycle4 NORMAL demo task_id = …", "正常链路关联 ID"],
            ["[IAM] OK … -> 200", "IAM 鉴权与执行成功"],
            ["[Agents/Mock] HTTP 500 … TaskPlan …（如有）", "文档助手 Mock 内部校验问题，与 IAM 网关演示无关"],
            ["正常演示完成 + agent_id / session_token / cert_id", "用于登录 Dashboard 或 Swagger"],
            ["Cycle4 ABNORMAL demo …", "异常链路 ID"],
            ["预期 403 / PERMISSION_DENIED", "拦截符合预期"],
            ["审计条数 … total=1", "拒绝已入账，可按 task_id 查询"],
            ["Done. Opened browser …", "已尝试打开 Dashboard"],
        ],
    )

    add_shot_block(
        doc,
        "图1 · 一键演示主窗口（健康检查与步骤）",
        "说明：演示脚本先完成 IAM 探活，再继续后续步骤，体现启动编排与安全就绪顺序。",
        "【答辩自用｜插入截图后删除本行】请将截图粘贴在上方「说明」段落后；"
        "画面须包含：标题 Agentrust one-click demo Cycle 4、OK http://127.0.0.1:8000/health，以及 Step 2～Step 5 的步骤标题行。",
    )

    add_shot_block(
        doc,
        "图2 · IAM 进程窗口（Uvicorn）",
        "说明：IAM 网关进程真实监听本机 8000，与探活 URL 一致。",
        "【答辩自用｜插入截图后删除本行】请将截图粘贴在上方「说明」段落后；"
        "画面须包含：窗口标题 Agentrust-IAM，以及 module=main:app、port=8000、Uvicorn running on http://127.0.0.1:8000。",
    )

    add_shot_block(
        doc,
        "图3 · 终端：正常演示输出",
        "说明：同一 task_id 下 IAM 多次执行成功，并打印可追溯标识与会话信息（可对 token 打码）。",
        "【答辩自用｜插入截图后删除本行】请将截图粘贴在上方「说明」段落后；"
        "画面须包含：Cycle4 NORMAL demo、[IAM] OK 多行、正常演示完成附近的 task_id、agent_id、session_token。",
    )

    add_shot_block(
        doc,
        "图4 · 终端：异常演示输出",
        "说明：企业 Agent 白名单与 IAM 能力校验两层拦截，拒绝行为进入审计。",
        "【答辩自用｜插入截图后删除本行】请将截图粘贴在上方「说明」段落后；"
        "画面须包含：预期 403、预期拒绝 HTTP 403、审计条数 total=1、异常场景的 task_id。",
    )

    # --- 5 ---
    add_heading(doc, "5. 步骤二：Dashboard（登录与任务链路）", level=1)
    doc.add_paragraph("浏览器访问：http://localhost:5173")
    doc.add_paragraph("登录（演示推荐）：")
    for line in [
        "在「演示 Session Token」中粘贴正常演示结束时打印的完整 session_token（JWT）。",
        "「Agent ID」须与同一段输出中的 agent_id 完全一致。",
        "点击登录；成功后进入仪表盘。",
    ]:
        doc.add_paragraph(line, style="List Number")
    doc.add_paragraph(
        "任务链路：左侧菜单进入「任务链路」，输入正常或异常演示的 task_id 查询；"
        "可查看各步允许/拒绝状态及「查看上下文」。"
    )

    add_shot_block(
        doc,
        "图5 · Dashboard 登录页",
        "说明：使用 IAM 签发的会话对接前端，以便加载审计与任务链路数据。",
        "【答辩自用｜插入截图后删除本行】请将截图粘贴在上方「说明」段落后；"
        "画面须包含：演示 Session Token 区域、Agent ID 已填写（Token 可部分打码）。",
    )

    add_shot_block(
        doc,
        "图6 · Dashboard 任务链路",
        "说明：按 task_id 聚合展示多步审计轨迹。",
        "【答辩自用｜插入截图后删除本行】请将截图粘贴在上方「说明」段落后；"
        "画面须包含：左侧任务链路选中、输入的 task_id、步骤列表中含 Agent/动作/允许或拒绝。",
    )

    add_shot_block(
        doc,
        "图7 · Dashboard 审计日志（可选）",
        "说明：审计列表与任务链路互为补充，可按任务筛选。",
        "【答辩自用｜插入截图后删除本行】请将截图粘贴在上方「说明」段落后；"
        "画面须包含：审计日志页及与演示 task_id 相关的记录。",
    )

    # --- 6 ---
    add_heading(doc, "6. 步骤三（可选）：Swagger 核验", level=1)
    doc.add_paragraph("访问：http://127.0.0.1:8000/docs")
    doc.add_paragraph(
        "Authorize 中使用 Bearer <同一 session_token>；调用 GET /api/v1/audit/trace/{task_id}，"
        "路径参数填控制台打印的 task_id。返回 JSON 即表示审计 API 与会话一致。"
    )

    add_shot_block(
        doc,
        "图8 · Swagger：Authorize 与 trace 响应",
        "说明：不依赖前端即可用 OpenAPI 复核链路数据。",
        "【答辩自用｜插入截图后删除本行】请将截图粘贴在上方「说明」段落后；"
        "画面须包含：Authorize 已配置、audit/trace 响应片段（含 trace 步骤概要）。",
    )

    # --- appendix ---
    add_heading(doc, "7. 附录：单独运行演示脚本", level=1)
    doc.add_paragraph("前提：IAM 及各 Agent 已按端口启动（仅测 IAM 时可只启 8000）。")
    add_code_block(
        doc,
        r"cd /d E:\yy_project\Agentrust\backend"
        "\n"
        r"venv\Scripts\activate"
        "\n"
        r"python scripts\demo_cycle4_normal.py"
        "\n"
        r"python scripts\demo_cycle4_abnormal.py",
    )
    doc.add_paragraph(
        "可选环境变量：IAM_BASE（默认 http://127.0.0.1:8000/api/v1）、DOC_HELPER_URL、ENTERPRISE_URL。"
    )

    add_heading(doc, "8. 附录：飞书 Scope 模板文件", level=1)
    doc.add_paragraph(
        "仓库文件 backend/scripts/feishu_app_scopes.cycle4.template.json 用于对照真实飞书开放平台权限勾选；"
        "当前演示以本地 Mock 与 IAM 为主。"
    )

    add_heading(doc, "9. 附录：常见问题", level=1)
    add_table(
        doc,
        ["现象", "说明"],
        [
            ["IAM 窗口路径/卷标语法错误", "请使用仓库内 run_demo.bat / spawn_uvicorn.bat 的启动方式"],
            ["登录成功仍停在登录页", "请使用当前版本前端：登录成功后整页跳转 /dashboard"],
            ["[Agents/Mock] 500 TaskPlan", "文档助手 Mock 校验问题，不影响 IAM 演示结论"],
            ["无 npm", "Dashboard 跳过时可改用 Swagger 与终端输出"],
        ],
    )

    add_heading(doc, "10. 演示结束如何关闭服务", level=1)
    doc.add_paragraph(
        "关闭标题为 Agentrust-IAM、Agent-Enterprise、Agent-ExternalSearch、Agent-DocHelper、"
        "Agentrust-Dashboard 的 CMD 窗口，或在对应窗口使用 Ctrl+C。"
    )

    doc.add_paragraph(
        "文档与仓库脚本 run_demo.bat、demo_cycle4_normal.py、demo_cycle4_abnormal.py 及 Dashboard 登录流程对齐。"
    )

    doc.save(out)
    try:
        import shutil

        shutil.copyfile(out, out_alias)
    except OSError:
        pass
    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()
